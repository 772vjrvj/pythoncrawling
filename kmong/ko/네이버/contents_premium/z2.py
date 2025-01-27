import os
import time
import psutil
import pyautogui
import pyperclip
from selenium.common.exceptions import WebDriverException
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium import webdriver
from openpyxl import load_workbook
from docx import Document
import os
import win32com.client

def delete_paragraph(paragraph):
    """문단 삭제 함수"""
    p_element = paragraph._element
    p_element.getparent().remove(p_element)
    p_element._p = p_element._element = None

def delete_text_between_multiple(doc_path, delete_ranges):
    # 문서 열기
    doc = Document(doc_path)

    paragraphs_to_remove = []
    deleting = False
    current_start = None
    current_end = None

    for paragraph in doc.paragraphs:
        if not deleting:
            # 각 구간의 시작 문구를 확인
            for start_text, end_text in delete_ranges:
                if start_text in paragraph.text:
                    deleting = True
                    current_start = start_text
                    current_end = end_text
                    paragraphs_to_remove.append(paragraph)
                    break  # 시작 문구 확인되면 다음 문구로 넘어감
        else:
            # 삭제 상태인 경우 현재 구간 처리
            paragraphs_to_remove.append(paragraph)
            if current_end in paragraph.text:
                deleting = False  # 현재 구간 종료
                current_start = None
                current_end = None

    # 삭제 대상 문단 제거
    for paragraph in paragraphs_to_remove:
        delete_paragraph(paragraph)

    # 기존 파일 덮어쓰기
    doc.save(doc_path)
    print(f"Updated document saved in-place: {doc_path}")


def close_chrome_processes():
    """모든 Chrome 프로세스를 종료합니다."""
    for proc in psutil.process_iter(['pid', 'name']):
        try:
            if 'chrome' in proc.info['name'].lower():
                proc.kill()  # Chrome 프로세스를 종료
        except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
            pass

def setup_driver():
    try:
        close_chrome_processes()

        chrome_options = Options()
        user_data_dir = f"C:\\Users\\{os.getlogin()}\\AppData\\Local\\Google\\Chrome\\User Data"
        profile = "Default"

        chrome_options.add_argument(f"user-data-dir={user_data_dir}")
        chrome_options.add_argument(f"profile-directory={profile}")
        chrome_options.add_argument("--disable-gpu")
        chrome_options.add_argument("--no-sandbox")
        chrome_options.add_argument("--disable-dev-shm-usage")
        chrome_options.add_argument("--disable-software-rasterizer")
        chrome_options.add_argument("--window-size=960,720")

        user_agent = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        chrome_options.add_argument(f'user-agent={user_agent}')
        chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
        chrome_options.add_experimental_option('useAutomationExtension', False)

        driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)

        script = '''
                Object.defineProperty(navigator, 'webdriver', { get: () => undefined });
                window.navigator.chrome = { runtime: {} };
                Object.defineProperty(navigator, 'languages', { get: () => ['en-US', 'en'] });
                Object.defineProperty(navigator, 'plugins', { get: () => [1, 2, 3, 4, 5] });
                Object.defineProperty(navigator, 'userAgent', { get: () => 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36' });
            '''
        driver.execute_cdp_cmd('Page.addScriptToEvaluateOnNewDocument', {'source': script})

        return driver
    except WebDriverException as e:
        print(f"Error setting up the WebDriver: {e}")
        return None

def perform_mouse_and_keyboard_actions():
    """마우스 클릭과 키보드 동작 수행."""
    try:
        # 첫 번째 클릭
        pyautogui.moveTo(1012, 83, duration=1)
        pyautogui.click()
        time.sleep(1)

        # 두 번째 클릭
        pyautogui.moveTo(599, 251, duration=1)
        pyautogui.click()
        time.sleep(1)

        # 세 번째 클릭
        pyautogui.moveTo(79, 186, duration=1)
        pyautogui.click()
        time.sleep(1)

        # Ctrl + A (전체 선택)
        pyautogui.hotkey('ctrl', 'a')
        time.sleep(1)

        # Ctrl + C (복사)
        pyautogui.hotkey('ctrl', 'c')
        time.sleep(1)

        # 클립보드 내용 반환
        return pyperclip.paste()

    except Exception as e:
        print(f"Error during mouse and keyboard actions: {e}")
        return None


def paste_to_docx_with_format(index):
    # 실행 경로를 기준으로 파일 저장 경로 생성
    current_path = os.getcwd()  # 현재 실행 경로 가져오기
    output_path = os.path.join(current_path, f"doc/output_{index}.docx")  # 파일 경로 설정

    # Word 애플리케이션 시작
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False  # Word 창 숨기기

    # 새 문서 생성
    doc = word.Documents.Add()

    # 클립보드에서 붙여넣기
    selection = word.Selection
    try:
        # 기본 붙여넣기 (서식 포함)
        wdPasteDefault = 0  # 하드코딩된 상수 값
        selection.PasteAndFormat(wdPasteDefault)
    except Exception as e:
        print(f"Error with PasteAndFormat: {e}")
        try:
            selection.Paste()  # 일반 붙여넣기
        except Exception as paste_error:
            print(f"Error with Paste: {paste_error}")

    # 문서 저장
    doc.SaveAs(output_path)
    doc.Close()
    word.Quit()
    print(f"Document saved to: {output_path}")
    return f"output_{index}.docx"


if __name__ == "__main__":
    excel_path = "비플랜 목차 테스트.xlsx"
    if not os.path.exists(excel_path):
        raise FileNotFoundError(f"엑셀 파일이 존재하지 않습니다: {excel_path}")

    workbook = load_workbook(excel_path)
    sheet = workbook.active

    urls = [row[0] for row in sheet.iter_rows(min_row=2, max_col=1, values_only=True)]

    doc = Document()
    doc.add_heading('비플랜 콘텐츠', level=1)

    driver = setup_driver()

    for index, url in enumerate(urls):
        try:
            driver.get(url)
            time.sleep(3)  # 페이지 로드 대기

            # 마우스 및 키보드 동작 수행 후 클립보드에서 데이터 가져오기
            content = perform_mouse_and_keyboard_actions()

            if content:
                output_path = paste_to_docx_with_format(index)

                delete_text_between_multiple(
                    doc_path=output_path,
                    delete_ranges=[
                        ("본문 바로가기", "해외 드랍쉬핑 실전판매의 정석"),  # 첫 번째 구간
                        ("해당 콘텐츠는 프리미엄 구독자", "NAVER Corp.")  # 두 번째 구간
                    ]
                )

        except Exception as e:
            print(f"Error processing URL {url}: {e}")

    driver.quit()

    print(f"Word 문서가 저장되었습니다")
