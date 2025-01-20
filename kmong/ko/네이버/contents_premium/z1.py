import os
import time
import psutil
import pyautogui
from selenium.common.exceptions import WebDriverException
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium import webdriver
from openpyxl import load_workbook
from docx import Document
import pyperclip

def close_chrome_processes():
    """모든 Chrome 프로세스를 종료합니다."""
    for proc in psutil.process_iter(['pid', 'name']):
        try:
            if 'chrome' in proc.info['name'].lower():
                proc.kill()  # Chrome 프로세스를 종료
        except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
            pass

def setup_driver():
    try:
        close_chrome_processes()

        chrome_options = Options()
        user_data_dir = f"C:\\Users\\{os.getlogin()}\\AppData\\Local\\Google\\Chrome\\User Data"
        profile = "Default"

        chrome_options.add_argument(f"user-data-dir={user_data_dir}")
        chrome_options.add_argument(f"profile-directory={profile}")
        chrome_options.add_argument("--disable-gpu")
        chrome_options.add_argument("--no-sandbox")
        chrome_options.add_argument("--disable-dev-shm-usage")
        chrome_options.add_argument("--disable-software-rasterizer")
        chrome_options.add_argument("--window-size=960,720")

        user_agent = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        chrome_options.add_argument(f'user-agent={user_agent}')
        chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
        chrome_options.add_experimental_option('useAutomationExtension', False)

        driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)

        script = '''
                Object.defineProperty(navigator, 'webdriver', { get: () => undefined });
                window.navigator.chrome = { runtime: {} };
                Object.defineProperty(navigator, 'languages', { get: () => ['en-US', 'en'] });
                Object.defineProperty(navigator, 'plugins', { get: () => [1, 2, 3, 4, 5] });
                Object.defineProperty(navigator, 'userAgent', { get: () => 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36' });
            '''
        driver.execute_cdp_cmd('Page.addScriptToEvaluateOnNewDocument', {'source': script})

        return driver
    except WebDriverException as e:
        print(f"Error setting up the WebDriver: {e}")
        return None

def perform_mouse_and_keyboard_actions():
    """마우스 클릭과 키보드 동작 수행."""
    try:
        # 첫 번째 클릭
        pyautogui.moveTo(1012, 83, duration=1)
        pyautogui.click()
        time.sleep(1)

        # 두 번째 클릭
        pyautogui.moveTo(599, 251, duration=1)
        pyautogui.click()
        time.sleep(1)

        # 세 번째 클릭
        pyautogui.moveTo(79, 186, duration=1)
        pyautogui.click()
        time.sleep(1)

        # Ctrl + A (전체 선택)
        pyautogui.hotkey('ctrl', 'a')
        time.sleep(1)

        # Ctrl + C (복사)
        pyautogui.hotkey('ctrl', 'c')
        time.sleep(1)

    except Exception as e:
        print(f"Error during mouse and keyboard actions: {e}")

if __name__ == "__main__":
    excel_path = "비플랜 목차 테스트.xlsx"
    if not os.path.exists(excel_path):
        raise FileNotFoundError(f"엑셀 파일이 존재하지 않습니다: {excel_path}")

    workbook = load_workbook(excel_path)
    sheet = workbook.active

    urls = [row[0] for row in sheet.iter_rows(min_row=2, max_col=1, values_only=True)]

    doc = Document()
    doc.add_heading('비플랜 콘텐츠', level=1)

    driver = setup_driver()

    for url in urls:
        try:
            driver.get(url)
            time.sleep(3)  # 페이지 로드 대기

            perform_mouse_and_keyboard_actions()

            # 클립보드에서 텍스트 가져오기
            content = pyperclip.paste()
            doc.add_paragraph(content)

        except Exception as e:
            print(f"Error processing URL {url}: {e}")

    driver.quit()

    output_path = "비플랜 테스트.docx"
    doc.save(output_path)
    print(f"Word 문서가 저장되었습니다: {output_path}")
