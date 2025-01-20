import os
import re
import requests
from bs4 import BeautifulSoup
import pandas as pd
import time
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_BREAK  # 추가

# 송창록.xlsx 파일 읽기
file_path = os.path.join(os.getcwd(), "송창록 목록.xlsx")
data = pd.read_excel(file_path)

# 데이터프레임을 객체 리스트로 변환
records = data.to_dict("records")

# 결과 리스트
result_list = []

# 요청 헤더 정의
headers = {
    "authority": "brunch.co.kr",
    "method": "GET",
    "scheme": "https",
    "accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7",
    "accept-encoding": "gzip, deflate, br, zstd",
    "accept-language": "ko-KR,ko;q=0.9,en-US;q=0.8,en;q=0.7",
    "cache-control": "max-age=0",
    "cookie": "__T_=1; __T_SECURE=1; b_s_a_l=1; __T_=1; __T_SECURE=1; ja_bt=7eF1DYRbLy8LCLcy:1737368443; ruuid=6bde92b3-5922-41d4-84ce-d4c7f38aee4b; _T_ANO=fKXUokVndt8mPqKw1KPX/5dHkCWKRRM+ig3CXehosu0exvczFXDsEnwP7LbFdkKKLJsKQx6GreddlgWyfBFd5Kq5vGOFlx7X0nDBxxJtgNF4lDdGe+EiP6w/ydj1z5o/iSzpJlJOXFIBZzwQVSdRGJlgu9yvDeG4JegyyJfOGl3joN7S0vTss6NfXAnOt5uIrJ4iEgCrGYgAXCj2KM8VuZc9dgHmI+U8jVnK82QNAd46TM3AYZqj9d2h6zaT4NyPgRN31zD09YxQJASGNEZC8hALX+qbXBsyLflZ0wQWA19r7Xk3CbFDnI1fwv07cQwgiBeBLyFj4X6f9EcExy0+yw==",
    "priority": "u=0, i",
    "sec-ch-ua": '"Google Chrome";v="131", "Chromium";v="131", "Not_A Brand";v="24"',
    "sec-ch-ua-mobile": "?0",
    "sec-ch-ua-platform": '"Windows"',
    "sec-fetch-dest": "document",
    "sec-fetch-mode": "navigate",
    "sec-fetch-site": "none",
    "sec-fetch-user": "?1",
    "upgrade-insecure-requests": "1",
    "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36"
}

# URL 요청 및 데이터 추출
for index, record in enumerate(records):
    url = record.get("URL")
    if not url:
        continue

    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()

        soup = BeautifulSoup(response.content, "html.parser")

        # 제목 추출
        title_tag = soup.find("h1", class_="cover_title")
        title = title_tag.text.strip() if title_tag else ""

        # 내용 추출
        content_div = soup.find("div", class_="wrap_body text_align_left")
        paragraphs = content_div.find_all("span") if content_div else []
        content = "\n".join(p.text.strip() for p in paragraphs if p.text.strip())

        obj = {
            "URL": url,
            "제목": title,
            "내용": content
        }

        # 결과 저장
        result_list.append(obj)

        print(obj)

        time.sleep(0.5)

    except Exception as e:
        print(f"Error processing URL {url}: {e}")

# Word 문서 생성
doc = Document()

for item in result_list:
    # 제목 추가
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(item['제목'])
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 내용 추가
    doc.add_paragraph(item['내용'])

    # 페이지 나누기
    page_break = doc.add_paragraph()
    page_break._element.getparent().remove(page_break._element)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)  # Ctrl+Enter


# Word 문서 저장
output_path = os.path.join(os.getcwd(), "송창록_결과.docx")
doc.save(output_path)

print(f"Word 문서가 생성되었습니다: {output_path}")
